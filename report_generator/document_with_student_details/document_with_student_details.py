
__author__ = "Tafadzwa Brian Motsi"

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Cm, Inches
from docx.shared import RGBColor
import random


class DocumentWithStudentDetails:

    def heading_of_the_report(self, font_name, student_name, grade_value, date):
        heading_names = ["THE LITTLE SCHOOL",
                         "PO.BOX 637",
                         "Mbabane, Pinevalley",
                         "Tell: 2411 1823",
                         "Cell: 76420 6321 / 7611 4551",
                         "Student: " + self.student_name(student_name),
                         "Grade: " + self.grade(grade_value),
                         "School re-opens next term: " + self.open_date(date)]

        document = Document()
        paragraph = document.sections[0].header.paragraphs[0]
        paragraph1 = document.add_paragraph("\n")
        count = 1
        for name in heading_names:

            if name is "THE LITTLE SCHOOL" and count < 6:
                run = paragraph.add_run(name)
                run.bold = True
                run.underline = True
                font = run.font
                font.name = font_name
                font.size = Pt(22)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            else:
                if count < 6:
                    run = paragraph.add_run("\n" + name)
                    font = run.font
                    font.name = font_name
                    font.size = Pt(15)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if count >= 6:
                if name != heading_names[-1]:
                    run = paragraph1.add_run(name + "\n")
                else:
                    run = paragraph1.add_run(name)
                font = run.font
                font.name = font_name
                font.size = Pt(15)
                paragraph1.alignment = WD_ALIGN_PARAGRAPH.LEFT

            count = count + 1

        return document

    def generate_documents(self, student_details_list, font_name, details_file_name, path_to_save_documents, date):
        student_dict = student_details_list[0]
        table_titles = ["Subject", "Tests", "Exam", "Term", "Remarks"]

        for key in student_dict:
            student_name = self.format_name(key)
            document = self.heading_of_the_report(font_name, student_name, details_file_name, date)

            second_paragraph = document.paragraphs[0].text

            for each_character in second_paragraph:
                pass
                if each_character == ":":
                    pass

            table = document.add_table((student_details_list[-1] + 3), 5, style="Table Grid")

            index_cell = 0

            for table_title in table_titles:
                table.rows[0].cells[index_cell].text = str(table_title)
                table.rows[0].cells[index_cell].paragraphs[0].runs[0].font.bold = True
                table.rows[0].cells[index_cell].paragraphs[0].runs[0].font.name = font_name
                table.rows[0].cells[index_cell].paragraphs[0].runs[0].font.size = Pt(15)
                table.rows[0].cells[index_cell].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

                index_cell = index_cell + 1

            index_cell = 0

            index_row = 1
            tests_totals = 0
            exams_totals = 0
            term_totals = 0

            totals_list = []
            averages_list = []

            term_totals_list = []

            count = 0
            for value in student_dict[key]:
                tests_totals = tests_totals + value[-3]
                exams_totals = exams_totals + value[-2]
                term_totals = term_totals + value[-1]
                term_totals_list.append(value[-1])
                count = count + 1

                for element in value:
                    table.rows[index_row].cells[index_cell].text = str(element)
                    table.rows[index_row].cells[index_cell].paragraphs[0].runs[0].font.name = font_name
                    table.rows[index_row].cells[index_cell].paragraphs[0].runs[0].font.size = Pt(14)
                    table.rows[index_row].cells[index_cell].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                    if str(element).isdigit():
                        if int(element) < 50:
                            table.rows[index_row].cells[index_cell].paragraphs[0].runs[0].font.color.rgb = \
                                RGBColor(255, 0, 0)
                        else:
                            pass
                    else:
                        pass

                    if index_cell < len(value)-1:
                        index_cell = index_cell + 1
                    else:
                        index_cell = 0
                        index_row = index_row + 1

            totals_list.append(tests_totals)
            totals_list.append(exams_totals)
            totals_list.append(term_totals)

            index_row = 1
            for term_mark in term_totals_list:
                if len(self.remarks(term_mark)) > 1:
                    table.rows[index_row].cells[-1].text = str(random.choice(self.remarks(term_mark)))
                else:
                    table.rows[index_row].cells[-1].text = str(self.remarks(term_mark)[0])

                table.rows[index_row].cells[-1].paragraphs[0].runs[0].font.name = font_name
                table.rows[index_row].cells[-1].paragraphs[0].runs[0].font.size = Pt(14)
                table.rows[index_row].cells[-1].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                index_row = index_row + 1

            averages_list.append(int(round(tests_totals/count)))
            averages_list.append(int(round(exams_totals/count)))
            averages_list.append(int(round(term_totals/count)))

            self.populate_table("Total", table, -2, 0, font_name)
            self.populate_table_1(totals_list, table, -2, 0, font_name)
            self.populate_table("Average", table, -1, 0, font_name)
            self.populate_table_1(averages_list, table, -1, 0, font_name)

            self.populate_table_2(table, term_mark, -1, -1, term_totals_list, count, font_name)

            self.populate_table_2(table, term_mark, -2, -1, term_totals_list, count, font_name)

            for _index in range(len(table.columns)):
                for cell in table.columns[_index].cells:
                    if _index == 0:
                        cell.width = Inches(2.0)
                        cell.height = Inches(3.0)
                    elif _index == 4:
                        cell.width = Inches(2.5)
                    else:
                        cell.width = Inches(1.0)

            for row in table.rows:
                row.height = Cm(0.8)

            self.closing_paragraph(document, font_name, averages_list[-1])
            document.save(path_to_save_documents + r"\\" + str(student_name) + ".docx")

    def closing_paragraph(self, document, font_name, term_average):
        scale_allocations = ["Merit", "First", "Second", "Third", "Fail"]
        closing_paragraph_element_list = ["Scale:\t75-100\t65-74\t\t50-64\t\t40-49\t\t0-39",
                                          scale_allocations,
                                          "\nComments:" + "." * 96 + '\n' + "." * 114 + ('\n' + "." * 114)*2,
                                          "\n"+"."*30 + '\t'*4 + "."*30 + " "*10 + "Class Teacher" + '\t'*5 +
                                          "Headmistress"
                                          ]

        index = 0
        paragraph = document.add_paragraph("\n")
        for element in closing_paragraph_element_list:
            index = index + 1
            if index != 2:
                run = paragraph.add_run(element + "\n")
                font = run.font
                font.name = font_name
                font.size = Pt(15)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            elif index == 2:
                for scale_allocation in scale_allocations:
                    if scale_allocation != scale_allocations[3]:
                        run_1 = paragraph.add_run("\t")
                        paragraph.add_run("\t")
                        run = paragraph.add_run(scale_allocation)

                    else:
                        paragraph.add_run("\t")
                        run = paragraph.add_run(scale_allocation)

                    font = run.font
                    font.name = font_name
                    font.size = Pt(15)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    if self.scale(term_average) == scale_allocation:
                        font.bold = True
                        font.underline = True

                    else:
                        if scale_allocation == scale_allocations[1]:
                            if self.scale(term_average) != scale_allocations[2]:
                                if self.scale(term_average) != scale_allocations[3]:
                                    if self.scale(term_average) != scale_allocations[4]:
                                        run_1.clear()
                                    else:
                                        pass
                                else:
                                    pass

                            else:
                                pass

                        elif scale_allocation == scale_allocations[4]:
                            if self.scale(term_average) == scale_allocations[3]:
                                run_1.clear()
                            else:
                                pass

                paragraph.add_run("\n")

    @staticmethod
    def scale(mark):
        if 0 <= mark <= 39:
            return "Fail"
        elif 40 <= mark <= 49:
            return "Third"
        elif 50 <= mark <= 64:
            return "Second"
        elif 65 <= mark <= 74:
            return "First"
        else:
            return "Merit"

    @staticmethod
    def remarks(mark):
        if 90 <= mark <= 100:
            return ["Excellent work", "Pleasing performance"]

        elif 80 <= mark <= 89:
            return ["Very good", "Well done"]

        elif 70 <= mark <= 79:
            return ["Good work", "Has the ability"]

        elif 60 <= mark <= 69:
            return ["Make an effort"]

        elif 50 <= mark <= 59:
            return ["Work hard"]
        else:
            return ["Work harder"]

    @staticmethod
    def student_name(name):
        return name

    @staticmethod
    def grade(name):
        digit = None
        for character in name:
            if str(character).isdigit():
                digit = character
                break
            else:
                digit = ""
        return digit

    @staticmethod
    def open_date(date):
        if date is not None:
            return date
        else:
            return ""

    @staticmethod
    def populate_table_1(totals_list, table, row_index, start_cell_index, font_name):
        for test_total in totals_list:
            index_cell = start_cell_index + 1
            table.rows[row_index].cells[index_cell].text = str(test_total)
            table.rows[row_index].cells[index_cell].paragraphs[0].runs[0].font.name = font_name
            table.rows[row_index].cells[index_cell].paragraphs[0].runs[0].font.size = Pt(14)
            table.rows[row_index].cells[index_cell].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
            table.rows[row_index].cells[index_cell].paragraphs[0].runs[0].font.bold = False
            start_cell_index = start_cell_index + 1

    @staticmethod
    def populate_table(label, table, row_index, index_cell, font_name):
        table.rows[row_index].cells[index_cell].text = label
        table.rows[row_index].cells[index_cell].paragraphs[0].runs[0].font.name = font_name
        table.rows[row_index].cells[index_cell].paragraphs[0].runs[0].font.bold = False
        table.rows[row_index].cells[index_cell].paragraphs[0].runs[0].font.size = Pt(14)
        table.rows[row_index].cells[index_cell].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

    def populate_table_2(self, table, term_mark, row_index, cell_index, totals_list, count, font_name):
        if len(self.remarks(term_mark)) > 1:
            table.rows[row_index].cells[cell_index].text = str(random.choice(self.remarks(sum(totals_list)/count)))
        else:
            table.rows[row_index].cells[cell_index].text = str(self.remarks(sum(totals_list)/count)[0])
        table.rows[row_index].cells[cell_index].paragraphs[0].runs[0].font.name = font_name
        table.rows[row_index].cells[cell_index].paragraphs[0].runs[0].font.bold = False
        table.rows[row_index].cells[cell_index].paragraphs[0].runs[0].font.size = Pt(14)
        table.rows[row_index].cells[cell_index].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

    @staticmethod
    def format_name(string):
        string = str(string).split(' ')
        index = 0
        for value in string:
            string[index] = value.capitalize()
            index = index + 1
        return ' '.join(str(x) for x in string)



